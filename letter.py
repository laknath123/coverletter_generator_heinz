# -*- coding: utf-8 -*-
"""
Created on Wed Dec 23 15:39:38 2020

@author: lakna
"""

#from reportlab.pdfgen.canvas import Canvas
from docx import Document
from job_descrip import enter_job_descrip
document = Document()

def cgenerator():
    
    
    full_name=input("What is your full name:? ")
    position=input("what is the position?: ")
    organization=input("What is the organization you are applying for?: ")
    jobsite=input("How did you come across this posting: ")
    enter_job_descrip()
    tasks=input("what type of tasks will you be doing at the job:? ")
    org_add=input("what is the address of the organization you are applying for: ?")
    sector=input("what specific thing are you passionate about relating to this job:? ")
    reason_1=input("What is the first reason you would be qualified for this job: ")
    reason_2=input("Include a second reason why you would be a good fit: ")
    reason_3=input("Include a third reason why you would be a good fit: ")
    para_3=input("use one of the things you mentioned earlier\n{}\n{}\n{}\n elaborate:".format(reason_1,reason_2,reason_3))
    
    org_add_split=org_add.split(',')
    
    address="2626 South Kingshighway Bvld\n Apartment 3E\n St. Louis-MO\n 63139"
    
    para1="""I am writing to express my strong interest in the {} position at {} that I came across through {}. 
    As you will see from my resume,my previous work experience and the rigorous data analytics curriculum at Heinz College at Carnegie Mellon has provided me strong 
    foundation in {}""".format(position,organization,jobsite,tasks)
    
    para2="""Combined with my passion for {} I think I would be a good for this
    {} at {} because {}""".format(sector,position,organization,reason_1)
    
    para3=para_3
    
    para4="""I would love the opportunity to discuss my qualifications with you in more detail. 
    I can be reached by phone or email, both listed on my resume. Thank you for your time and consideration; I look forward to speaking with you soon."""
    
    since="Sincerely"
    
    
    
    
    document.add_heading(full_name+" Cover Letter", 0)
    
    p = document.add_paragraph(address)
    p = document.add_paragraph(org_add)
    p = document.add_paragraph("Dear")
    p = document.add_paragraph(para1)
    p = document.add_paragraph(para2)
    p = document.add_paragraph(para3)
    p = document.add_paragraph(para4)
    p = document.add_paragraph(since)
    p = document.add_paragraph(full_name)
    
    document.save('coverletter.docx')
    
    #canvas = Canvas("coverletter.pdf")
    #canvas.drawString(1, 1,para1)
    #canvas.save()


cgenerator()